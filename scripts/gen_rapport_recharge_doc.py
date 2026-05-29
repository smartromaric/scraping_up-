#!/usr/bin/env python3
"""Génère rapport_recharge_chauffeurs.docx."""
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

DRIVERS = [
    ("BIAGNET FELICIEN DODO", "AA-369-SJ-01"),
    ("ISSAU YUSSUF AYO", "2025-55238 WWW-01"),
    ("Diallo issouf", "AA922AX01"),
    ("BOGUI", "AA-894-TF-01"),
    ("DAKOURI KOUDOU GERVAIS JUNIOR", "AB-928-AK"),
    ("FOFANA GOGBE YOUSSOUF", "AA-102-HC-01"),
    ("Yacouba Konaté", "AA-623-BB"),
    ("Diomande lancina", "AA116VF01"),
    ("KONAN KOFFI ENOC", "AA156AJ01"),
    ("CISSE OUSMANE", "AA-247-VK-01"),
    ("Tre bi tizie Didier", "AA318RV01"),
    ("KAMBIRE BONTIKO KOKO JOEL", "977KR01"),
    ("TRAORE AMARA FELIX", "AA-013-XQ-01"),
    ("Keita", "AA-791-BV"),
    ("SERY RICHMOND ROMEO", "AA-693-JJ"),
]

OUT = Path(__file__).resolve().parent.parent / "output" / "partner_automation" / "rapport_recharge_chauffeurs.docx"


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Rapport de recharge — chauffeurs partenaire", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Objet : Liste des conducteurs ayant reçu une recharge portefeuille")
    doc.add_paragraph(f"Date du rapport : {datetime.now().strftime('%d/%m/%Y')}")
    doc.add_paragraph("Montant unitaire : 2 000 FCFA par conducteur")
    doc.add_paragraph()

    doc.add_heading("Synthèse", level=1)
    doc.add_paragraph(f"Nombre de conducteurs rechargés : {len(DRIVERS)}")
    doc.add_paragraph("Statut véhicule : Approuvé")
    doc.add_paragraph(
        "Critère : conducteur assigné à un véhicule approuvé sur la flotte partenaire"
    )
    doc.add_paragraph()

    doc.add_heading("Liste des chauffeurs rechargés", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "N°"
    hdr[1].text = "Nom du conducteur"
    hdr[2].text = "Numéro de plaque"
    for i, (name, plate) in enumerate(DRIVERS, 1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = name
        row[2].text = plate

    doc.add_paragraph()
    doc.add_heading("Commentaire", level=1)
    doc.add_paragraph(
        "Les recharges ont été effectuées pour les conducteurs listés ci-dessus, "
        "tous rattachés à un véhicule au statut Approuvé sur la flotte partenaire. "
        "Ce document constitue le relevé des bénéficiaires de l'opération de recharge portefeuille."
    )
    doc.add_paragraph()
    note = doc.add_paragraph(
        "Document établi à partir du relevé flotte partenaire (conducteur + plaque)."
    )
    note.italic = True

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
