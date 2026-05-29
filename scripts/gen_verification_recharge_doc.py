#!/usr/bin/env python3
"""Génère verification_rapports_recharge.docx depuis rapport_recharge_verification_state.json."""
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

SCRIPT_DIR = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(SCRIPT_DIR))

JSON_PATH = SCRIPT_DIR / "output" / "partner_automation" / "rapport_recharge_verification_state.json"
OUT = SCRIPT_DIR / "output" / "partner_automation" / "verification_rapports_recharge.docx"


def main() -> None:
    with JSON_PATH.open(encoding="utf-8") as f:
        rows = json.load(f)

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    t = doc.add_heading("Vérification recharges — rapports vs state.json", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Pour chaque nom des rapports PDF (20/05 et 21/05), comparaison avec "
        "transfer_2000_done dans state.json."
    )
    doc.add_paragraph()

    by_report: dict[str, list] = {}
    for r in rows:
        rap = r.get("rapport", "?")
        key = f"{rap}|{r.get('nom_rapport')}|{r.get('plaque_rapport')}"
        if key not in {f"{x.get('rapport')}|{x.get('nom_rapport')}|{x.get('plaque_rapport')}" for x in by_report.get(rap, [])}:
            by_report.setdefault(rap, []).append(r)

    for rap, items in by_report.items():
        doc.add_heading(rap, level=1)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        h = table.rows[0].cells
        h[0].text = "Nom (rapport)"
        h[1].text = "Plaque"
        h[2].text = "Statut"
        h[3].text = "Partenaire"
        h[4].text = "Détail"
        for it in items:
            c = table.add_row().cells
            c[0].text = it.get("nom_rapport", "")
            c[1].text = it.get("plaque_rapport", "")
            c[2].text = it.get("statut", "")
            c[3].text = str(it.get("partenaire", ""))
            c[4].text = it.get("detail", "")
        doc.add_paragraph()

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
