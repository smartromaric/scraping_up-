# -*- coding: utf-8 -*-
"""Génère le document Word des cas de test répartition (communs aux deux équipes)."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = Path(__file__).parent / "Cas_de_test_repartition_UPJUNOO.docx"

EQUIPES = [
    "Equipe Angré CHU",
    "Equipe Palmeraie et Saint Viateur",
]

# (numéro, situation, données, résultats attendus)
CAS_COMMUNS = [
    (
        "Test 1",
        "Affectation au chauffeur le plus proche",
        "Un client lance une commande (zone Angré CHU ou Palmeraie / Saint Viateur). "
        "5 chauffeurs dans la zone : A à 300 m ; B à 800 m ; C à 1,5 km ; D à 2 km ; E à 3 km. "
        "Tous disponibles, wallet valide, compte actif.",
        "Le système propose la commande au chauffeur A en priorité (proximité + éligibilité).",
    ),
    (
        "Test 2",
        "Refus par le premier chauffeur",
        "Le chauffeur le plus proche refuse la commande.",
        "Transfert automatique au chauffeur suivant disponible ; pas de blocage ; "
        "client toujours en attente ; nouvelle proposition envoyée rapidement.",
    ),
    (
        "Test 3",
        "Chauffeur sans solde wallet suffisant",
        "Un chauffeur est proche mais son wallet est insuffisant.",
        "Le système ne lui attribue pas la commande ; proposition à un autre chauffeur "
        "disponible avec wallet valide.",
    ),
    (
        "Test 4",
        "Commandes simultanées",
        "Plusieurs clients (ex. 5) lancent une commande au même moment dans la même zone.",
        "Le répartiteur évite les doublons ; n'affecte pas un même chauffeur à deux commandes ; "
        "affecte chaque commande à un chauffeur disponible ; respecte la proximité géographique.",
    ),
    (
        "Test 5",
        "Saturation locale",
        "Plusieurs clients (ex. 8) commandent alors que seulement 5 engins sont proches "
        "(Palmeraie, Angré CHU ou autre zone saturée).",
        "Affectation des engins disponibles en priorité ; clients excédentaires en attente ; "
        "élargissement progressif du rayon ; proposition d'engins plus éloignés si nécessaire.",
    ),
    (
        "Test 6",
        "Chauffeur qui accepte puis annule",
        "Un chauffeur accepte une commande puis annule avant ou pendant la prise en charge.",
        "La commande revient dans le circuit de répartition ; le client reçoit une notification claire.",
    ),
    (
        "Test 7",
        "Perte de connexion chauffeur",
        "Un chauffeur reçoit une proposition mais perd la connexion ou ne répond pas.",
        "Après expiration du délai d'acceptation, la commande est proposée à un autre chauffeur.",
    ),
    (
        "Test 8",
        "Mauvaise géolocalisation",
        "Un chauffeur apparaît loin sur la carte alors qu'il est physiquement proche du client.",
        "Anomalie visible dans l'admin ; vérifier : précision GPS, autorisation localisation, "
        "actualisation de position, délai de mise à jour carte.",
    ),
    (
        "Test 9",
        "Chauffeur occupé (en course)",
        "Le chauffeur le plus proche est déjà en course ; d'autres chauffeurs sont disponibles.",
        "La commande n'est pas proposée au chauffeur occupé ; proposition au suivant "
        "selon proximité et disponibilité.",
    ),
    (
        "Test 10",
        "Timeout sans réponse",
        "Le chauffeur le plus proche reçoit la proposition mais ne répond pas dans le délai configuré.",
        "Après expiration : proposition au chauffeur suivant ; client informé ; aucune double affectation.",
    ),
    (
        "Test 11",
        "Chauffeur en pause ou hors ligne",
        "Des chauffeurs proches sont visibles sur la carte mais en pause ou hors ligne.",
        "Seuls les chauffeurs au statut Disponible sont éligibles ; les autres sont exclus.",
    ),
    (
        "Test 12",
        "Égalité de distance (ex æquo)",
        "Deux chauffeurs à la même distance du point de prise en charge, tous deux éligibles.",
        "Règle de tie-break cohérente et documentée ; un seul chauffeur reçoit la proposition en premier.",
    ),
    (
        "Test 13",
        "Mise à jour de position en cours de recherche",
        "La position du chauffeur le plus proche change pendant la recherche (s'éloigne ou se rapproche).",
        "Le répartiteur recalcule sur position fraîche ; bascule vers un autre chauffeur si pertinent, "
        "sans bloquer le client.",
    ),
    (
        "Test 14",
        "Respect du périmètre d'équipe / zone",
        "Un chauffeur proche est hors périmètre de l'équipe ou de la zone de la commande.",
        "Priorité aux chauffeurs rattachés à la zone ; pas d'affectation cross-zone non autorisée.",
    ),
    (
        "Test 15",
        "Annulation client avant acceptation",
        "Le client annule pendant la recherche chauffeur (proposition déjà envoyée, pas encore acceptée).",
        "Proposition retirée chez le chauffeur ; commande clôturée ; pas de sollicitation inutile des suivants.",
    ),
    (
        "Test 16",
        "Annulation client après acceptation",
        "Le chauffeur a accepté ; le client annule avant prise en charge effective.",
        "Course annulée ; chauffeur redevient disponible ; notification chauffeur ; journal cohérent.",
    ),
    (
        "Test 17",
        "Compte ou documents non validés",
        "Un chauffeur proche a un compte bloqué (documents en attente, KYC incomplet).",
        "Chauffeur exclu de la répartition ; proposition à un chauffeur actif ; trace visible en admin.",
    ),
    (
        "Test 18",
        "Wallet insuffisant puis recharge",
        "Chauffeur proche sous le seuil wallet, recharge réussie pendant la fenêtre de recherche.",
        "Tant que wallet invalide : pas d'affectation ; après recharge validée : éligibilité selon règle métier.",
    ),
    (
        "Test 19",
        "Adresse ou point de prise en charge imprécis",
        "Le client place le pin loin de sa position GPS réelle.",
        "Répartition basée sur le point de commande ; alerte possible adresse approximative ; "
        "correction possible par le superviseur.",
    ),
    (
        "Test 20",
        "Aucun chauffeur dans le rayon initial",
        "0 engin disponible dans le rayon court ; engins disponibles plus loin.",
        "Client en attente ; élargissement progressif du rayon ; estimation délai ; pas de boucle infinie.",
    ),
    (
        "Test 21",
        "Ordre de la file d'attente (FIFO)",
        "Saturation : plus de clients que d'engins disponibles.",
        "Les premières commandes chronologiques sont affectées en priorité ; "
        "les suivantes restent en attente dans l'ordre d'arrivée.",
    ),
    (
        "Test 22",
        "Délai maximum d'attente client",
        "Client en attente au-delà du seuil configuré malgré élargissement du rayon.",
        "Notification client ; proposition d'engin plus éloigné ou escalade superviseur ; "
        "statut visible en admin.",
    ),
    (
        "Test 23",
        "Chaîne de refus / annulations multiples",
        "Premier chauffeur annule, second refuse, troisième disponible.",
        "Tentatives successives sans blocage ; client informé à chaque étape ; "
        "affectation finale ou échec contrôlé selon limite de tentatives.",
    ),
    (
        "Test 24",
        "Ne pas reproposer au même chauffeur",
        "Un chauffeur a refusé ou a expiré (timeout) sur une commande.",
        "La même commande ne revient pas automatiquement à ce chauffeur "
        "(sauf réassignation manuelle superviseur).",
    ),
    (
        "Test 25",
        "Réassignation manuelle par le superviseur",
        "Le répartiteur automatique ne trouve pas d'engin éligible ; le superviseur force un chauffeur.",
        "Action tracée (qui, quand, motif) ; client notifié ; chauffeur verrouillé sur la course.",
    ),
    (
        "Test 26",
        "Échec de notification push",
        "Chauffeur éligible mais notification push non reçue (token invalide, app en arrière-plan).",
        "Fallback in-app / SMS si configuré ; sinon timeout puis proposition au chauffeur suivant.",
    ),
    (
        "Test 27",
        "Application fermée ou tuée",
        "Le chauffeur ferme l'application après réception de la proposition sans réponse serveur.",
        "Comportement identique au timeout ; pas de course fantôme assignée.",
    ),
    (
        "Test 28",
        "Position GPS obsolète",
        "Dernière mise à jour de position du chauffeur au-delà du seuil (ex. > 10 min).",
        "Chauffeur exclu de la répartition automatique ; flag position obsolète en admin.",
    ),
    (
        "Test 29",
        "Client injoignable au point de prise en charge",
        "Chauffeur arrivé sur place, client ne répond pas.",
        "Procédure d'attente puis annulation motivée ; chauffeur libéré ; règles frais / pénalité appliquées.",
    ),
    (
        "Test 30",
        "Double commande du même client",
        "Le client lance deux commandes en peu de temps au même endroit.",
        "Détection doublon ; une seule course active ; message client explicite.",
    ),
    (
        "Test 31",
        "Session multiple sur un même compte chauffeur",
        "Même compte connecté sur deux appareils simultanément.",
        "Une seule réponse d'acceptation valide ; pas de double acceptation concurrente.",
    ),
    (
        "Test 32",
        "Audit et traçabilité post-course",
        "Une course est menée à bien après répartition automatique.",
        "Admin affiche : chauffeur retenu, distances au moment T, refus/timeouts, "
        "durée attente client, statut wallet, horodatages GPS.",
    ),
    (
        "Test 33",
        "Renfort inter-zones (si politique activée)",
        "Zone saturée (ex. Palmeraie) ; engins disponibles dans une autre zone (ex. Angré CHU).",
        "Si politique désactivée : attente locale uniquement. Si activée : renfort tracé, "
        "priorité aux clients en attente la plus longue.",
    ),
    (
        "Test 34",
        "Sous-zones au sein d'une même équipe",
        "Commandes simultanées Palmeraie et Saint Viateur avec engins répartis par sous-zone.",
        "Priorité aux engins de la sous-zone de la commande ; pas de vol d'engins sans règle explicite.",
    ),
    (
        "Test 35",
        "Charge concurrente (test de robustesse)",
        "Nombre élevé de commandes simultanées sur les deux équipes.",
        "Aucune double affectation ; latence acceptable ; logs sans erreur critique ; "
        "cohérence des files d'attente.",
    ),
]


def set_cell_shading(cell, hex_color: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_header_row(table, labels, bg="016D71"):
    row = table.rows[0]
    for i, text in enumerate(labels):
        cell = row.cells[i]
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
        set_cell_shading(cell, bg)


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    title = doc.add_heading("Cas de test — Répartition des commandes", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("UPJUNOO — Tests communs aux équipes terrain")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(1, 109, 113)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Date : {date.today().strftime('%d/%m/%Y')}\n").font.size = Pt(11)
    meta.add_run("Périmètre : ").bold = True
    meta.add_run(", ".join(EQUIPES))

    doc.add_paragraph()
    intro = doc.add_paragraph(
        "Ce document regroupe l'ensemble des cas de test communs applicables aux deux équipes "
        "de supervision. Chaque scénario doit être exécuté dans au moins une zone de référence "
        "(Angré CHU, Palmeraie ou Saint Viateur) et reproduit si besoin sur l'autre équipe pour "
        "valider la cohérence du répartiteur."
    )
    intro.paragraph_format.space_after = Pt(12)

    doc.add_heading("Équipes concernées", level=2)
    for eq in EQUIPES:
        doc.add_paragraph(eq, style="List Bullet")

    doc.add_heading("Grille des cas de test communs", level=2)

    headers = [
        "Équipe(s)",
        "N° test",
        "Situation",
        "Données / contexte",
        "Résultats attendus",
        "Zone(s) de validation",
        "Superviseur",
        "Statut",
        "Observations",
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    add_header_row(table, headers)

    zones_suggest = {
        1: "Angré CHU",
        2: "Angré CHU / Palmeraie",
        3: "Les deux",
        4: "Les deux",
        5: "Palmeraie (prioritaire), Angré CHU",
        6: "Les deux",
        7: "Les deux",
        8: "Les deux",
        9: "Les deux",
        10: "Les deux",
        11: "Les deux",
        12: "Angré CHU",
        13: "Les deux",
        14: "Les deux",
        15: "Les deux",
        16: "Les deux",
        17: "Les deux",
        18: "Les deux",
        19: "Les deux",
        20: "Palmeraie / Angré CHU",
        21: "Palmeraie (saturation)",
        22: "Palmeraie",
        23: "Les deux",
        24: "Les deux",
        25: "Les deux",
        26: "Les deux",
        27: "Les deux",
        28: "Les deux",
        29: "Les deux",
        30: "Les deux",
        31: "Les deux",
        32: "Les deux",
        33: "Palmeraie ↔ Angré CHU",
        34: "Palmeraie + Saint Viateur",
        35: "Les deux (charge)",
    }

    for numero, situation, donnees, resultats in CAS_COMMUNS:
        test_num = int(numero.split()[-1])
        row = table.add_row()
        values = [
            "Commun",
            numero,
            situation,
            donnees,
            resultats,
            zones_suggest.get(test_num, "Les deux"),
            "",
            "À exécuter",
            "",
        ]
        for i, text in enumerate(values):
            cell = row.cells[i]
            cell.text = text
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
        if test_num % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, "F0F9FA")

    doc.add_page_break()
    doc.add_heading("Légende — Statut d'exécution", level=2)
    for label, desc in [
        ("À exécuter", "Test non encore réalisé."),
        ("En cours", "Test démarré, résultat non finalisé."),
        ("OK", "Comportement conforme aux résultats attendus."),
        ("KO", "Écart constaté — anomaly à traiter."),
        ("N/A", "Non applicable dans l'environnement de test."),
    ]:
        p = doc.add_paragraph()
        p.add_run(f"{label} : ").bold = True
        p.add_run(desc)

    doc.add_heading("Critères de réussite globaux", level=2)
    criteres = [
        "Aucune double affectation d'un même chauffeur sur deux commandes actives.",
        "Respect de la proximité géographique parmi les chauffeurs éligibles.",
        "Exclusion systématique des chauffeurs non disponibles, wallet insuffisant ou compte bloqué.",
        "Transparence admin : refus, timeouts, GPS, wallet et historique de répartition consultables.",
        "Expérience client : pas de blocage silencieux, notifications en cas d'annulation ou de délai.",
    ]
    for c in criteres:
        doc.add_paragraph(c, style="List Bullet")

    return doc


def main():
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Document créé : {OUTPUT}")


if __name__ == "__main__":
    main()
